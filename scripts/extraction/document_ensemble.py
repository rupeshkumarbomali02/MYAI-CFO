#!/usr/bin/env python3
import argparse, json, os, re, sys, tempfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

# Keep local extraction bounded so a heavy PDF does not spawn an uncontrolled
# pile of OCR/layout-model workers on a CFO laptop.
MAX_WORKERS = max(1, min(2, int(os.getenv('MYAI_EXTRACTOR_WORKERS', '2'))))
for k in ('OMP_NUM_THREADS','OPENBLAS_NUM_THREADS','MKL_NUM_THREADS'):
    os.environ.setdefault(k, '2')
os.environ.setdefault('TOKENIZERS_PARALLELISM', 'false')

CRO_SCALE = 10_000_000.0

CURRENCY_SYMBOLS = {'₹':'INR','$':'USD','£':'GBP','€':'EUR','¥':'JPY','₫':'VND','₩':'KRW','R$':'BRL','A$':'AUD','C$':'CAD'}

def detect_money_metadata(text):
    t=text or ''
    # Strong document-level indicators must outrank incidental currency mentions
    # elsewhere in a global annual report (for example SAR appearing in a later
    # subsidiary disclosure must never redefine a US-dollar consolidated report).
    strong_patterns=[
        (r'(?i)\(\s*dollars\s+in\s+(million|millions|billion|billions|thousand|thousands)\s*\)', 'USD'),
        (r'(?i)\bdollars\s+in\s+(million|millions|billion|billions|thousand|thousands)\b', 'USD'),
        (r'(?i)\(\s*us\s+dollars\s+in\s+(million|millions|billion|billions|thousand|thousands)\s*\)', 'USD'),
        (r'(?i)\$\s*in\s+(million|millions|billion|billions|thousand|thousands)\b', 'USD'),
        (r'(?i)\(\s*in\s+(million|millions|billion|billions|thousand|thousands)\s*,?\s*except\s+per\s+share', 'USD'),
    ]
    scale_map={'mn':'million','bn':'billion','m':'million','k':'thousand'}
    for pat,ccy in strong_patterns:
        m=re.search(pat,t)
        if m:
            scale_word=m.group(1).lower().rstrip('s')
            return ccy, scale_word, f'{ccy} {scale_word}'
    # Prefer table/header labels such as 'Amount (USD million)' before generic
    # currency mentions. This prevents a header like 'USD' from collapsing a
    # million-scale financial statement back to units.
    labelled=re.search(r'(?i)\b(?:amounts?|amount|figures?|values?|currency)\s*\(\s*(INR|USD|GBP|EUR|JPY|CNY|CAD|AUD|SGD|HKD|AED|IDR|ZAR|BRL|MXN|SAR|CHF|NOK|SEK|DKK|NZD)\s+(crore|million|billion|thousand|lakh|mn|bn|m|k)\s*\)', t)
    if labelled:
        ccy=labelled.group(1).upper(); scale=scale_map.get(labelled.group(2).lower(), labelled.group(2).lower())
        return ccy, scale, f'{ccy} {scale}'
    # Prefer explicit document-level phrases such as 'All amounts in INR million'.
    explicit=re.search(r'(?:all\s+amounts?|amounts?|figures?|values?)\s+(?:are\s+)?(?:stated\s+)?(?:in|denominated\s+in)\s+\(?\s*(INR|USD|GBP|EUR|JPY|CNY|CAD|AUD|SGD|HKD|AED|IDR|ZAR|BRL|MXN|SAR|CHF|NOK|SEK|DKK|NZD)\s*(crore|million|billion|thousand|lakh|mn|bn|m|k)?', t, re.I)
    if explicit:
        ccy=explicit.group(1).upper(); scale=(explicit.group(2) or '').lower()
        scale_word=scale_map.get(scale,scale or 'units')
        return ccy, scale_word, f'{ccy} {scale_word}' if scale_word!='units' else ccy
    # Bare currency codes in narrative text are not document-level currency evidence.
    # Require an adjacent amount/scale marker to avoid false AUD/INR/etc. detections from prose, indexes or URLs.
    m=re.search(r'(?P<ccy>INR|USD|GBP|EUR|JPY|CNY|CAD|AUD|SGD|HKD|AED|IDR|ZAR|BRL|MXN|SAR|CHF|NOK|SEK|DKK|NZD)\s+(?P<scale>crore|million|billion|thousand|lakh|mn|bn|m|k)\b', t, re.I)
    if m:
        ccy=m.group('ccy').upper(); scale=m.group('scale').lower()
        scale_word=scale_map.get(scale,scale)
        return ccy, scale_word, f'{ccy} {scale_word}'
    for sym,ccy in sorted(CURRENCY_SYMBOLS.items(), key=lambda x:-len(x[0])):
        if sym in t:
            m2=re.search(r'\b(crore|million|billion|thousand|lakh|mn|bn|m|k)\b',t,re.I)
            scale_word=scale_map.get((m2.group(1).lower() if m2 else ''),'') if m2 else 'units'
            return ccy, scale_word or 'units', f'{ccy} {scale_word or "units"}'
    return None,None,None

def apply_document_scale(value, raw, scale):
    if value is None or not scale or scale == 'units':
        return value
    low=str(raw or '').lower()
    if re.search(r'\b(?:crore|cr|million|billion|thousand|lakh|mn|bn|m|k)\b', low):
        return value
    return value * {'crore':1e7,'million':1e6,'billion':1e9,'thousand':1e3,'lakh':1e5}.get(scale,1.0)

CONCEPT_PATTERNS = {
    'revenue': [r'^revenue from operations\b', r'^(?:total\s+)?revenues?\b'],
    'cogs': [r'^(?:cost of goods sold|cost of sales|cogs)\b'],
    'gross_profit': [r'^gross profit\b'],
    'operating_income': [r'^(?:operating income|profit from operations|operating profit)\b'],
    'ebitda': [r'ebitda'],
    'net_income': [r'^(?:profit for the year|profit attributable to|net income|net profit)\b'],
    'equity': [r'^(?:total\s+)?equity\b', r'^shareholders?[’\']?\s+equity\b', r'^owners?\s+equity\b'],
    'current_debt': [r'^(?:current debt|current portion of debt and finance leases)\b'],
    'long_term_debt': [r'^debt and finance leases\s*,?\s*net of current portion\b', r'^long[- ]term debt\b'],
    'cash': [r'cash and cash equivalents', r'cash equivalents', r'^cash$'],
    'current_assets': [r'^total current assets\b', r'^current assets(?:\s|$)'],
    'assets': [r'^(?:total\s+)?assets\b'],
    'receivables': [r'^(?:trade receivables|accounts receivable|receivables)\b'],
    'inventory': [r'inventor(?:y|ies)'],
    'current_liabilities': [r'^total current liabilities\b', r'^current liabilities(?:\s|$)'],
    'payables': [r'^(?:trade payables|accounts payable|payables)\b'],
    'debt': [r'^(?:total debt|debt|borrowings)\b'],
    'liabilities': [r'^(?:total\s+)?liabilities\b'],
    'capex': [r'^(?:capital expenditure|capital expenditures|purchase of property)\b'],
    'operating_cash_flow': [r'^(?:operating cash flow|net cash (?:flow|inflow|outflow) from operating activities)\b'],
    'depreciation_amortization': [r'depreciation\s*(?:and|&)\s*amortisation', r'depreciation\s*(?:and|&)\s*amortization', r'^depreciation\s*&\s*amortization\b'],
    'interest_expense': [r'^(?:finance costs|interest expense)\b'],
}

STATEMENT_HINTS = {
    'income': re.compile(r'(results of operations|statement of profit and loss|statement of operations|income statement|profit and loss|consolidated statements? of operations)', re.I),
    'balance': re.compile(r'(consolidated\s+balance\s+sheets?|balance sheets?|balance sheet|statement of financial position)', re.I),
    'cashflow': re.compile(r'(cash flows?|cash flow statement|statement of cash flows)', re.I),
}

YEAR_RE = re.compile(r'(?P<y1>20\d{2})\s*[-/]\s*(?P<y2>\d{2,4})')
FISCAL_YEAR_RE = re.compile(r'(?i)\b(?:fiscal\s+year|FY)\s*[:\-]?\s*(20\d{2})\b')
DATE_YEAR_RE = re.compile(r'(?:31(?:st|nd|rd|th)?\s+March|March\s+31(?:st|nd|rd|th)?|31(?:st|nd|rd|th)?\s+December|December\s+31(?:st|nd|rd|th)?)[,\s]+(?P<y>20\d{2})', re.I)
NUMBER_RE = re.compile(r'(?<![A-Za-z])(?:\(?-?\d{1,3}(?:,\d{2,3})*(?:\.\d+)?\)?)(?:\s*(?:million|billion|m|bn|crore|cr))?(?![A-Za-z])', re.I)


def clean_text(text):
    return re.sub(r'\s+', ' ', text or '').strip()


def normalize_num(raw):
    if raw is None:
        return None
    s = str(raw).strip().replace('₹','').replace('$','').replace('£','').replace('€','')
    neg = s.startswith('(') and s.endswith(')')
    s = s.strip('()').replace(',', '').strip()
    mult = 1.0
    low = s.lower()
    if low.endswith('crore') or re.search(r'\bcr$', low):
        mult = CRO_SCALE
        s = re.sub(r'\s*(?:crore|cr)$', '', s, flags=re.I)
    elif low.endswith('billion') or low.endswith('bn'):
        mult = 1e9
        s = re.sub(r'\s*(?:billion|bn)$', '', s, flags=re.I)
    elif low.endswith('million') or low.endswith('m'):
        mult = 1e6
        s = re.sub(r'\s*(?:million|m)$', '', s, flags=re.I)
    try:
        v = float(s) * mult
        return -v if neg else v
    except Exception:
        return None


def filename_document_year(path):
    try:
        name=Path(path).stem
        matches=list(re.finditer(r'(?<!\d)(20\d{2})(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])(?!\d)', name))
        for m in reversed(matches):
            year,month,day=int(m.group(1)),int(m.group(2)),int(m.group(3))
            # Filing filenames such as tsla-20251231.htm and msft-20250630.htm
            # encode the filing/reporting period. A valid YYYYMMDD is therefore
            # a stronger period signal than incidental future years in document text.
            if month>=1 and day>=1:
                return year
    except Exception:
        pass
    return None

def filename_comparative_year(path):
    """Return the latest explicit comparative year encoded in a fixture/report filename."""
    try:
        stem=Path(path).stem if path else ""
        years=[int(y) for y in re.findall(r"(?<!\d)(20\d{2})(?!\d)", stem)]
        years=sorted(set(years))
        return years[-1] if len(years)>=2 else None
    except Exception:
        return None

def detect_document_year(text, path=None):
    t=text or ''
    # Strongest signal: an annual-report / filing filename that encodes a year-end,
    # such as tsla-20251231.htm. This prevents future disclosure years in inline XBRL
    # from becoming the document period.
    filename_year=filename_document_year(path) if path else None
    if filename_year:
        return filename_year
    filename_range_year=filename_comparative_year(path) if path else None
    if filename_range_year:
        return filename_range_year
    # Next, prefer explicit statement period language near the beginning of the filing.
    head=t[:50000]
    patterns=[
        r'(?i)(?:as\s+of|year\s+ended|for\s+the\s+year\s+ended|for\s+the\s+fiscal\s+year\s+ended)[^\n]{0,180}?(?:december|31st|12/31|12\s*\-\s*31)[^\n]{0,60}?\b(20\d{2})\b',
        r'(?i)for\s+the\s+fiscal\s+year\s+ended[^\n]{0,160}?\b(20\d{2})\b',
        r'(?i)for\s+the\s+year\s+ended[^\n]{0,160}?\b(20\d{2})\b',
        r'(?i)fiscal\s+year\s*[:\-]\s*(20\d{2})',
        r'(?i)document\s+fiscal\s+year\s*[:\-]\s*(20\d{2})',
    ]
    for pat in patterns:
        m=re.search(pat,head)
        if m:
            return int(m.group(1))
    # Explicit FY headings are still useful, but choose the earliest reliable report
    # heading in the opening portion rather than the maximum year anywhere in the file.
    fiscal_headings=[int(m.group(1)) for m in FISCAL_YEAR_RE.finditer(head)]
    if fiscal_headings:
        # Comparative statements commonly contain one explicit heading per period.
        # The document period is the latest reported period, not the first comparative column.
        return max(set(fiscal_headings))
    years=[int(x) for x in re.findall(r'\b20\d{2}\b',head[:20000])]
    return years[0] if years else None


def period_candidates(text, default_year=None):
    years = []
    t = text or ''
    for m in YEAR_RE.finditer(t):
        y1 = int(m.group('y1')); y2 = int(m.group('y2'))
        y2 = y2 + 2000 if y2 < 100 else y2
        years.extend([y1, y2])
    for m in DATE_YEAR_RE.finditer(t):
        years.append(int(m.group('y')))
    for m in FISCAL_YEAR_RE.finditer(t):
        years.append(int(m.group(1)))
    # Table headers are frequently emitted as a bare sequence (`2025 2024 2023`)
    # or as consecutive standalone year lines. Capture both forms.
    standalone=[]
    for line in t.splitlines()[:120]:
        hdr=[int(x) for x in re.findall(r'(?<!\d)(20\d{2})(?!\d)', line)]
        if 2 <= len(set(hdr)) <= 4 and len(hdr) <= 4:
            years.extend(hdr); break
        if re.fullmatch(r'\s*20\d{2}\s*', line):
            standalone.append(int(line.strip()))
    if len(set(standalone)) >= 2:
        years.extend(standalone[:4])
    years = sorted(set(years), reverse=True)
    if default_year and default_year not in years:
        years.insert(0, default_year)
    return years[:4]

def row_fiscal_year(rows, idx, default_year=None):
    """Return the nearest explicit fiscal/report year heading for a row.

    Comparative financial statements frequently render the fiscal-year heading on
    its own line/page block. Using a single document-level year for every row would
    collapse FY2024/FY2025 facts into one year and silently lose the comparative
    period.
    """
    for j in range(idx, max(-1, idx-18), -1):
        line=rows[j] if j < len(rows) else ''
        m=FISCAL_YEAR_RE.search(line)
        if m:
            return int(m.group(1))
    return int(default_year) if default_year else None


def statement_context(text):
    low = text.lower()
    hits=[]
    for name,rx in STATEMENT_HINTS.items():
        m=rx.search(text)
        if m: hits.append((m.start(),m.end(),name))
    if not hits:
        low=low if 'low' in locals() else (text or '').lower()
        if re.search(r'\b(?:equity and liabilities|current liabilities|total current liabilities|total liabilities)\b', low): return 'balance'
        if re.search(r'\b(?:revenue from operations|profit for the year|total income|profit before tax)\b', low): return 'income'
        if re.search(r'\b(?:net cash flow from operating activities|cash and cash equivalents.*beginning|cash flows?)\b', low): return 'cashflow'
        return 'unknown'
    # Use the first statement heading in the page block.
    hits.sort()
    return hits[0][2]


def row_statement_context(rows, idx, fallback='unknown'):
    """Infer statement type from the nearest preceding statement heading.
    A single PDF page may contain P&L followed by balance sheet content, so a
    page-wide context is unsafe for financial fact extraction.
    """
    window_start=max(0, idx-24)
    segment=rows[window_start:idx+1]
    best=None
    for j,line in enumerate(segment):
        for name,rx in STATEMENT_HINTS.items():
            if rx.search(line):
                best=(j,name)
    return best[1] if best else fallback

def candidate_rows(text):
    lines=[re.sub(r'\s+',' ',x).strip() for x in (text or '').splitlines()]
    return [x for x in lines if x]


def primary_statement_pages(pages):
    # Parent consolidated statements are normally the first main income statement
    # and first balance sheet. Later pages often contain subsidiary statements.
    selected=[]
    for p in pages:
        tx=p.get('text','')
        low=tx.lower()
        if 'consolidated' in low and any(rx.search(tx) for rx in STATEMENT_HINTS.values()):
            selected.append(p)
        elif selected and statement_context(tx) != 'unknown' and len(selected)<12:
            selected.append(p)
    return selected[:16]


def facts_from_page(page, doc_year, filename):
    text=page.get('text','') or ''
    context=statement_context(text)
    if context=='unknown':
        return []
    years=period_candidates(text, doc_year)
    # Need two comparative years for normal statements. For a page with a single
    # current date, only use the document year and infer prior year from header where present.
    if not years and doc_year: years=[doc_year]
    rows=candidate_rows(text)
    out=[]
    for idx,line in enumerate(rows):
        low=line.lower()
        row_context=row_statement_context(rows, idx, context)
        if not low: continue
        concept=None
        for key,pats in CONCEPT_PATTERNS.items():
            if any(re.search(p, low, re.I) for p in pats):
                concept=key; break
        if concept is None: continue
        if concept == 'liabilities' and re.search(r'\bliabilit(?:y|ies)\s+(?:and|&)\s+equity\b', low, re.I):
            continue
        # Enforce statement-context discipline so a cash-flow delta, MD&A narrative,
        # note number, or subsidiary disclosure cannot become a headline financial fact.
        context_requirements={
            'cash':'balance','current_assets':'balance','assets':'balance','receivables':'balance',
            'inventory':'balance','current_liabilities':'balance','payables':'balance','debt':'balance','current_debt':'balance','long_term_debt':'balance','liabilities':'balance',
            'revenue':'income','cogs':'income','gross_profit':'income','operating_income':'income','ebitda':'income','net_income':'income','equity':'balance',
            'capex':'cashflow','operating_cash_flow':'cashflow','depreciation_amortization':'cashflow',
        }
        required=context_requirements.get(concept)
        if required and row_context != required:
            continue
        if row_context in ('income','balance','cashflow') and ('auditor' in low[:80] or 'key audit matter' in low[:80]):
            continue
        row_year=row_fiscal_year(rows, idx, doc_year)
        # Extract values tied to the row label. Financial statements legitimately contain
        # small values (including 1-2 digit figures), so never discard them based on size.
        # If a layout engine separates the label from its value, inspect only the immediately
        # following numeric line and stop if another financial concept begins.
        nums=[m.group(0) for m in NUMBER_RE.finditer(line)]
        source=line
        if not nums:
            collected=[]; collected_lines=[]
            max_lines=max(3,len(years)+2)
            for j in range(1,max_lines+1):
                if idx+j >= len(rows): break
                nxt=rows[idx+j]; nl=nxt.lower().strip()
                if any(any(re.search(pat,nl,re.I) for pat in pats) for pats in CONCEPT_PATTERNS.values()):
                    break
                if re.search(r'^(?:year ended|dollars in|as a percentage|change)\b',nl,re.I):
                    break
                if re.fullmatch(r'20\d{2}',nl):
                    continue
                n=[m.group(0) for m in NUMBER_RE.finditer(nxt)]
                if not n:
                    if collected: break
                    continue
                collected.extend(n); collected_lines.append(nxt)
                if len(collected) >= max(1,min(len(years),4)): break
            if collected:
                nums=collected[:4]; source=line+' | '+' | '.join(collected_lines)
        # When multiple numbers are present on the same row, a leading tiny number can be a
        # note reference. Only remove that leading token when more than two numeric tokens exist.
        # A financial-statement table may put a small note reference before the values.
        # Remove it only when the subsequent financial numbers are materially larger;
        # never discard legitimate two-digit financial values such as 96/88/80.
        if len(nums)>2 and re.fullmatch(r'\d{1,2}', nums[0].strip()):
            tail=[normalize_num(x) for x in nums[1:]]
            first=normalize_num(nums[0])
            if first is not None and tail and all(v is not None for v in tail) and min(abs(v) for v in tail) >= max(100.0,abs(first)*10):
                nums=nums[1:]
        if not nums: continue
        for pos,raw in enumerate(nums[:4]):
            value=normalize_num(raw)
            if value is None: continue
            # Multi-column financial tables bind values by header order. A three-column
            # statement (FY2025/FY2024/FY2023) must never collapse all numbers into FY2025.
            if len(nums)>1 and pos < len(years):
                fiscal=years[pos]
            else:
                fiscal=row_year if row_year is not None else (years[pos] if pos < len(years) else (doc_year if pos==0 else None))
            if fiscal is None: continue
            detected_currency, detected_scale, detected_unit = detect_money_metadata(text)
            unit = detected_unit
            currency = detected_currency
            absolute_scale={'crore':1e7,'million':1e6,'billion':1e9,'thousand':1e3,'lakh':1e5}.get(detected_scale or 'units',1.0)
            absolute_value=value*absolute_scale if value is not None else None
            if concept in {'revenue','cogs','gross_profit','operating_income','ebitda','net_income','cash','current_assets','assets','receivables','inventory','current_liabilities','payables','debt','current_debt','long_term_debt','liabilities','equity','capex','operating_cash_flow','depreciation_amortization','interest_expense'}:
                aggregateRole='reported-aggregate' if (re.search(r'\btotal\b',low,re.I) or re.match(r'^(?:current assets|current liabilities|assets|liabilities|equity)\b',low,re.I)) and not re.search(r'\bliabilit(?:y|ies)\s+(?:and|&)\s+equity\b',low,re.I) else 'source-line'
                out.append({
                    # V46 remediation contract: normalizedValue is the source numeric value
                    # exactly as presented (e.g. 94,827). Scale remains explicit metadata;
                    # absoluteValue/baseValue are available only for controlled compatibility
                    # calculations and must never be used for executive display.
                    'concept':concept,'sourceLabel':line.split('|')[0].strip(),'aggregateRole':aggregateRole,'rawValue':raw,'normalizedValue':value,'absoluteValue':absolute_value,'baseValue':absolute_value,'unit':unit or detected_unit or 'document unit',
                    'currency':currency,'scale':detected_scale or 'units','sourceUnitText':detected_unit,'fiscalYear':int(fiscal),'periodEnd':f'31 March {fiscal}' if 'March' in text else None,
                    'sourcePage':page.get('page'),'evidenceText':source[:1800],
                    'extractionMethod':page.get('extractor','unknown'),'statementContext':row_context,
                    'filename':filename,'confidence':0.98 if any(x in low for x in ('total ','revenue from operations','profit for the year')) else 0.93,
                    'systemVerified':False,'validated':False,'rowBinding':'exact-label-row'
                })
    return out


def run_pymupdf(path):
    import fitz
    doc=fitz.open(path)
    pages=[]
    for i,p in enumerate(doc): pages.append({'page':i+1,'text':p.get_text('text') or '' ,'extractor':'pymupdf'})
    return {'name':'PyMuPDF','method':'pymupdf','pages':pages,'tables':[]}


def run_pypdf(path):
    """Independent pure-Python PDF text fallback for PDFs whose MuPDF text layer is incomplete."""
    from pypdf import PdfReader
    reader=PdfReader(path)
    pages=[]
    for i,page in enumerate(reader.pages, start=1):
        text=page.extract_text() or ''
        pages.append({'page':i,'text':text,'extractor':'pypdf','tables':[]})
    return {'name':'pypdf','method':'pypdf','pages':pages,'tables':[]}

def run_tesseract_ocr(path):
    """Optional OCR fallback for image-only/scanned financial PDFs. Bounded to the first 40 pages."""
    import shutil
    if not shutil.which('tesseract'):
        raise RuntimeError('tesseract binary not installed')
    import fitz
    import pytesseract
    from PIL import Image
    doc=fitz.open(path)
    pages=[]
    limit=min(len(doc), int(os.getenv('MYAI_OCR_MAX_PAGES','40')))
    for i in range(limit):
        page=doc.load_page(i)
        pix=page.get_pixmap(matrix=fitz.Matrix(1.5,1.5), alpha=False)
        img=Image.frombytes('RGB',[pix.width,pix.height],pix.samples)
        text=pytesseract.image_to_string(img) or ''
        pages.append({'page':i+1,'text':text,'extractor':'tesseract-ocr','tables':[]})
    return {'name':'Tesseract OCR','method':'tesseract-ocr','pages':pages,'tables':[],'ocrPages':len(pages)}

def run_pdfplumber(path):
    import pdfplumber
    pages=[]
    with pdfplumber.open(path) as pdf:
        for i,p in enumerate(pdf.pages):
            text=p.extract_text() or ''
            tables=[]
            try:
                for t in p.extract_tables() or []:
                    tables.append(t)
            except Exception: pass
            pages.append({'page':i+1,'text':text,'extractor':'pdfplumber','tables':tables})
    return {'name':'pdfplumber','method':'pdfplumber','pages':pages,'tables':[t for p in pages for t in p.get('tables',[])]}


def run_docling(path):
    from docling.document_converter import DocumentConverter
    converter=DocumentConverter()
    result=converter.convert(path)
    doc=result.document
    text=doc.export_to_markdown()
    pages=[]
    for i,chunk in enumerate(text.split('\n\f') if '\f' in text else [text],1):
        pages.append({'page':i,'text':chunk,'extractor':'docling'})
    tables=[]
    try:
        for t in doc.tables:
            try: tables.append(t.export_to_dataframe().to_dict(orient='records'))
            except Exception: pass
    except Exception: pass
    return {'name':'Docling','method':'docling','pages':pages,'tables':tables}


def run_marker(path):
    # Optional: only invoked if a CLI is actually available.
    import shutil, subprocess
    cli=shutil.which('marker_single')
    if not cli: raise RuntimeError('marker_single not installed')
    with tempfile.TemporaryDirectory() as td:
        cp=subprocess.run([cli,path,'--output_dir',td],capture_output=True,text=True)
        if cp.returncode: raise RuntimeError(cp.stderr[-1000:])
        md=list(Path(td).rglob('*.md'))
        text=md[0].read_text(encoding='utf-8',errors='ignore') if md else cp.stdout
    return {'name':'Marker','method':'marker','pages':[{'page':1,'text':text,'extractor':'marker'}],'tables':[]}


def run_mineru(path):
    import shutil, subprocess
    cli=shutil.which('mineru')
    if not cli: raise RuntimeError('mineru not installed')
    with tempfile.TemporaryDirectory() as td:
        cp=subprocess.run([cli,'-p',path,'-o',td],capture_output=True,text=True)
        if cp.returncode: raise RuntimeError(cp.stderr[-1000:])
        md=list(Path(td).rglob('*.md'))
        text='\n'.join(f.read_text(encoding='utf-8',errors='ignore') for f in md)
    return {'name':'MinerU','method':'mineru','pages':[{'page':1,'text':text,'extractor':'mineru'}],'tables':[]}




def run_pdftable(path):
    """Optional independent table extractor using CycloneBoy/pdf_table.
    Never becomes mandatory: unavailable/failed PdfTable is recorded and existing
    pdfplumber/PyMuPDF evidence remains authoritative.
    """
    import shutil, subprocess, json, tempfile
    exe=shutil.which(os.getenv('MYAI_CFO_PDFTABLE_EXE','pdftable'))
    if not exe:
        return {'name':'PdfTable','method':'pdftable','pages':[],'tables':[],'error':'pdftable CLI not installed','optional':True}
    outdir=tempfile.mkdtemp(prefix='myai-pdftable-')
    output=os.path.join(outdir,'result.json')
    adapter=Path(__file__).resolve().parent.parent/'pdf'/'pdftable_adapter.py'
    try:
        cp=subprocess.run([sys.executable,str(adapter),'--input',str(path),'--output_dir',outdir,'--pages',os.getenv('MYAI_PDFTABLE_PAGES','all'),'--timeout',os.getenv('MYAI_PDFTABLE_TIMEOUT','240'),'--output',output],capture_output=True,text=True,timeout=int(os.getenv('MYAI_PDFTABLE_TIMEOUT','240'))+20)
        if os.path.exists(output):
            result=json.loads(Path(output).read_text(encoding='utf-8'))
        else:
            result={'status':'ERROR','tables':[],'pages':[],'error':cp.stderr[-2000:]}
        return {'name':'PdfTable','method':'pdftable','pages':result.get('pages',[]),'tables':result.get('tables',[]),'error':result.get('error') or (result.get('stderr') if result.get('status') not in ('PASS','NOT_INSTALLED') else None),'optional':True,'status':result.get('status'),'config':result.get('config',{})}
    except Exception as e:
        return {'name':'PdfTable','method':'pdftable','pages':[],'tables':[],'error':str(e),'optional':True}

def run_nvidia_retriever(path):
    """Optional NVIDIA NeMo Retriever / nv-ingest path.

    Uses the current Python API when installed. It is deliberately optional so
    the local-first stack still works when NVIDIA services/packages are absent.
    """
    try:
        from nv_ingest_client.client.interface import Ingestor
    except Exception as exc:
        raise RuntimeError(f'NVIDIA NeMo Retriever Python client not installed: {exc}')
    method=os.getenv('MYAI_NVIDIA_EXTRACT_METHOD','pdfium_hybrid').strip() or 'pdfium_hybrid'
    ingestor=(Ingestor().files(path).extract(
        extract_text=True,
        text_depth='page',
        extract_tables=True,
        extract_charts=True,
        extract_infographics=True,
        extract_images=False,
        extract_method=method,
    ))
    result=ingestor.ingest(show_progress=False,return_failures=True,return_traces=True)
    results=result[0] if isinstance(result,tuple) else result
    failures=result[1] if isinstance(result,tuple) and len(result)>1 else []
    if not results:
        raise RuntimeError(f'NVIDIA NeMo Retriever returned no document result; failures={failures!r}')

    def stringify(obj):
        try:
            return str(obj)
        except Exception:
            return ''

    def collect_text(obj,out):
        if obj is None:return
        if isinstance(obj,str):
            if len(obj)>20: out.append(obj)
            return
        if isinstance(obj,dict):
            for k,v in obj.items():
                lk=str(k).lower()
                if lk in {'text','content','text_content','markdown','page_text'} and isinstance(v,str):
                    if len(v)>20: out.append(v)
                else: collect_text(v,out)
            return
        if isinstance(obj,(list,tuple)):
            for v in obj: collect_text(v,out)
            return
        # dataclass / client result objects often expose dict-like fields via vars()
        try:
            d=vars(obj)
            if d:
                collect_text(d,out); return
        except Exception: pass
        t=stringify(obj)
        if t and len(t)>20: out.append(t)

    texts=[]; collect_text(results,texts)
    # De-duplicate while preserving order and cap pathological repr output.
    seen=set(); pages=[]
    for i,t in enumerate(texts,1):
        t=clean_text(t)
        if not t or t in seen: continue
        seen.add(t); pages.append({'page':i,'text':t,'extractor':'nvidia_nemo_retriever'})
    return {
        'name':'NVIDIA NeMo Retriever',
        'method':'nvidia-nemo-retriever',
        'pages':pages,
        'tables':[],
        'traces': result[2] if isinstance(result,tuple) and len(result)>2 else None,
        'failures': failures,
        'extractMethod':method,
    }

def _html_grid(table):
    """Expand HTML rowspan/colspan into a semantic rectangular grid."""
    grid=[]; pending={}
    for r_idx,tr in enumerate(table.find_all('tr')):
        while len(grid)<=r_idx: grid.append([])
        row=grid[r_idx]; col=0
        for cell in tr.find_all(['th','td']):
            while (r_idx,col) in pending: row.append(pending[(r_idx,col)]); col+=1
            text=clean_text(cell.get_text(' ',strip=True))
            try: colspan=max(1,int(cell.get('colspan') or 1))
            except Exception: colspan=1
            try: rowspan=max(1,int(cell.get('rowspan') or 1))
            except Exception: rowspan=1
            for rr in range(r_idx,r_idx+rowspan):
                while len(grid)<=rr: grid.append([])
                while len(grid[rr])<col: grid[rr].append('')
                for cc in range(col,col+colspan):
                    while len(grid[rr])<=cc: grid[rr].append('')
                    grid[rr][cc]=text
                    if rr>r_idx: pending[(rr,cc)]=text
            col+=colspan
        while (r_idx,col) in pending: row.append(pending[(r_idx,col)]); col+=1
    mx=max((len(r) for r in grid),default=0)
    for r in grid: r.extend(['']*(mx-len(r)))
    return grid

def _html_year_anchors(rows):
    best={}
    for row in rows[:15]:
        cur={}
        for i,cell in enumerate(row):
            m=re.fullmatch(r'\s*(20\d{2})\s*',str(cell or ''))
            if m: cur[int(m.group(1))]=i
        if len(cur)>=2 and len(cur)>len(best): best=cur
    return dict(sorted(best.items(),key=lambda kv:kv[1]))

def _html_years(rows): return list(_html_year_anchors(rows).keys())[:4]

def _html_concept(label):
    low=clean_text(label).lower()
    if re.search(r'liabilit(?:y|ies)\s+(?:and|&)\s+equity',low): return None
    checks=[
        ('current_assets',r'^(?:total\s+)?current assets$'),('current_liabilities',r'^(?:total\s+)?current liabilities$'),('assets',r'^(?:total\s+)?assets$'),('liabilities',r'^(?:total\s+)?liabilities$'),
        ('equity',r'^(?:total\s+)?equity$|^total stockholders[’\']? equity$|^shareholders[’\']? equity$'),
        ('current_debt',r'^(?:current debt|current portion of debt and finance leases)$'),('long_term_debt',r'^debt and finance leases,?\s*net of current portion$|^long[- ]term debt$'),('debt',r'^(?:total debt|total borrowings)$'),
        ('revenue',r'^(?:total\s+)?revenues?$|^revenue from operations$|^value of sales & services \(revenue\)$'),('gross_profit',r'^gross profit$'),('operating_income',r'^(?:income from operations|operating income)$'),
        ('net_income',r'^(?:net income|profit for the year)$'),('cash',r'^(?:cash and cash equivalents|cash equivalents)$'),('receivables',r'^(?:accounts receivable,? net|trade receivables|accounts receivable)$'),
        ('inventory',r'^(?:inventory|inventories)$'),('payables',r'^(?:accounts payable|trade payables)$'),('cogs',r'^(?:total cost of revenues|cost of revenues|cost of goods sold|cost of sales)$'),
        ('operating_cash_flow',r'^net cash provided by operating activities$'),('capex',r'^(?:capital expenditures|purchase of property, plant and equipment)$')]
    for c,rx in checks:
        if re.search(rx,low): return c
    return None

def _html_numeric(cell):
    v=normalize_num(cell)
    return v if re.fullmatch(r'[$£€₹]?\(?-?[\d,]+(?:\.\d+)?\)?',str(cell or '').strip()) else None

def run_html(path):
    from bs4 import BeautifulSoup
    raw=Path(path).read_text(errors='ignore'); soup=BeautifulSoup(raw,'html.parser')
    for tag in soup(['script','style','noscript']): tag.decompose()
    tables=[]; facts=[]; images=[]; statement_hint='unknown'; full_text=soup.get_text('\n')
    for h in soup.find_all(['h1','h2','h3','h4','title']):
        ht=clean_text(h.get_text(' ',strip=True)).lower()
        if re.search(r'consolidated.*operations|statement.*operations|profit.*loss|income statement',ht): statement_hint='income'
        elif re.search(r'consolidated.*balance|balance sheet|financial position',ht): statement_hint='balance'
        elif re.search(r'cash flows?|statement of cash flows',ht): statement_hint='cashflow'
    for ii,img in enumerate(soup.find_all('img'),1):
        src=clean_text(img.get('src',''))
        if src: images.append({'index':ii,'page':1,'src':src,'alt':clean_text(img.get('alt','')),'kind':'remote-image'})
    last_years=[]
    for ti,table in enumerate(soup.find_all('table'),1):
        rows=_html_grid(table)
        if not rows: continue
        anchors=_html_year_anchors(rows); years=list(anchors.keys())[:4]; last_years=years or last_years; mx=max(map(len,rows))
        year_header_index=0
        for rr,row in enumerate(rows[:15]):
            hits=sum(1 for cell in row if re.fullmatch(r'\s*20\d{2}\s*',str(cell or '')))
            if hits>=2: year_header_index=rr; break
        header_rows=rows[:year_header_index+1] if year_header_index>=0 else [rows[0]]
        tables.append({'page':1,'index':ti,'title':f'HTML financial table {ti}','rows':len(rows),'columns':mx,'headers':header_rows[-1] if header_rows else rows[0],'headerRows':header_rows,'yearHeaderRowIndex':year_header_index,'fiscalYears':years,'yearColumns':anchors,'rowsData':rows,'tsv':'\n'.join('\t'.join(r) for r in rows)})
        for r in rows:
            concept=_html_concept(r[0] if r else '')
            if not concept or not anchors: continue
            for fy,col in anchors.items():
                candidates=[]
                for cc in range(col,min(mx,col+5)):
                    cell=r[cc] if cc<len(r) else ''
                    nv=_html_numeric(cell)
                    if nv is not None and nv not in years: candidates.append((cc,cell,nv))
                if not candidates: continue
                cc,rawv,v=min(candidates,key=lambda x:x[0])
                label=r[0]
                facts.append({'concept':concept,'sourceLabel':label,'aggregateRole':'reported-aggregate' if (re.search(r'\btotal\b',label,re.I) or re.match(r'^(?:current assets|current liabilities|assets|liabilities|equity)$',label,re.I)) else 'source-line','rawValue':rawv,'normalizedValue':v,'absoluteValue':v,'baseValue':v,'unit':'document unit','currency':None,'scale':'units','sourceUnitText':None,'fiscalYear':fy,'sourcePage':1,'evidenceText':' | '.join(r)[:1800],'extractionMethod':'html-table-structured','statementContext':statement_hint,'confidence':0.995,'systemVerified':False,'validated':False,'rowBinding':'exact-html-table-cell','tableIndex':ti,'sourceColumn':cc,'sourceYearColumn':col})
    fy=filename_document_year(path) or (max(last_years) if last_years else detect_document_year(full_text,path))
    return {'name':'BeautifulSoup HTML tables','method':'html-table-structured','pages':[{'page':1,'text':full_text,'extractor':'html-table-structured'}],'tables':tables,'images':images,'structuredFacts':facts,'documentFiscalYear':fy,'comparativeFiscalYears':sorted(set(last_years),reverse=True),'extractionQuality':{'grade':'high' if facts else ('rich' if tables else 'text'),'comparativeFiscalYears':sorted(set(last_years),reverse=True),'comparativeColumnCount':len(set(last_years)),'tableCount':len(tables),'imageCount':len(images),'financialDocument':True,'factCount':len(facts),'ragFeedExtractor':'html-table-structured'}}

COMPREHENSIVE_CANONICAL_ALIASES = {
    'D&A':'depreciation_amortization','Depreciation':'depreciation','Equity':'equity','ARPU':'avg_revenue_per_customer','Cash & Cash Equivalents':'cash','Capital Expenditures':'capex','Diluted EPS Reported':'eps','SG&A Expenses':'sga_expenses',
    'Expansion ARR':'expansion','Contraction ARR':'contraction','Churn ARR':'churn_arr',
    'Starting ARR':'starting_arr','Net New ARR':'net_new_arr','Sales & Marketing Spend':'sales_marketing_spend',
    'Starting Revenue':'starting_revenue','Lost Revenue':'lost_revenue','Starting Customers':'starting_customers',
    'Customers Lost':'customers_lost','MRR Per Customer':'mrr_per_customer','Current Quarter Revenue':'current_quarter_revenue',
    'Prior Quarter Revenue':'prior_quarter_revenue','Prior Quarter Sales & Marketing Spend':'prior_quarter_sales_marketing_spend',
    'Net Cash Burn':'net_cash_burn','Prior Revenue':'prior_revenue','One-off Adjustments':'one_off_adjustments',
    'Fixed Costs':'fixed_costs','Price Per Unit':'price_per_unit','Variable Cost Per Unit':'variable_cost_per_unit',
    'Actual Sales':'actual_sales','Break-even Sales':'break_even_sales'
}

def _excel_canonical(label):
    s=str(label or '').strip()
    if s in COMPREHENSIVE_CANONICAL_ALIASES: return COMPREHENSIVE_CANONICAL_ALIASES[s]
    s=s.lower().replace('&',' and ').replace('’',"'")
    s=re.sub(r'[^a-z0-9]+','_',s).strip('_')
    return {'cash_cash_equivalents':'cash','total_assets':'assets','total_liabilities':'liabilities','total_debt':'debt','accounts_receivable':'receivables','accounts_payable':'payables','shareholders_equity':'equity','tax_rate':'tax_rate'}.get(s,s)

def extract_comprehensive_workbook(path):
    import openpyxl
    wb=openpyxl.load_workbook(path,data_only=True,read_only=True)
    if 'Comprehensive Inputs' not in wb.sheetnames: return []
    ws=wb['Comprehensive Inputs']; rows=list(ws.iter_rows(values_only=True))
    if not rows: return []
    header=[str(v).strip() if v is not None else '' for v in rows[0]]
    year_cols={}
    for i,h in enumerate(header):
        if re.fullmatch(r'20\d{2}',h): year_cols[int(h)]=i
    facts=[]
    for rr in rows[1:]:
        if not rr or rr[0] in (None,''): continue
        label=str(rr[0]).strip(); concept=_excel_canonical(label)
        if not concept: continue
        unit=str(rr[1] or '').strip() if len(rr)>1 else ''
        for fy,col in year_cols.items():
            if col>=len(rr) or rr[col] is None or rr[col]=='': continue
            val=rr[col]
            try:
                n=float(val)
            except Exception:
                continue
            # Workbook percentages are stored as decimal fractions; mark them explicitly.
            fact_unit='decimal' if unit=='%' and abs(n)<=1 else unit or 'USD mm'
            facts.append({'concept':concept,'sourceLabel':label,'aggregateRole':'analytical-input','rawValue':str(val),'normalizedValue':n,'absoluteValue':n,'baseValue':n,'unit':fact_unit,'currency':'USD','scale':'units','sourceUnitText':unit or None,'fiscalYear':fy,'sourcePage':1,'evidenceText':f'Comprehensive Inputs | {label} | {fy}: {val}','extractionMethod':'openpyxl-comprehensive-inputs','statementContext':'analytical-input','confidence':1.0,'systemVerified':True,'validated':False,'rowBinding':'excel-input-row-year-column','verificationMethod':'workbook-explicit-input'} )
    # Populate supporting analytical inputs that are directly derivable from the workbook's explicit rows.
    # These remain source-traceable as workbook-derived rather than model-invented.
    by={};
    for f in facts: by.setdefault((f['concept'],int(f['fiscalYear'])),f)
    years=sorted({int(f['fiscalYear']) for f in facts})
    def add_derived(concept, fy, val, label, formula):
        if (concept,fy) in by or val is None: return
        facts.append({'concept':concept,'sourceLabel':label,'aggregateRole':'workbook-derived-input','rawValue':str(val),'normalizedValue':float(val),'absoluteValue':float(val),'baseValue':float(val),'unit':'decimal' if concept in ('revenue_growth_rate','organic_growth') else 'USD mm','currency':'USD','scale':'units','sourceUnitText':None,'fiscalYear':fy,'sourcePage':1,'evidenceText':f'Workbook-derived support | {formula} | {fy}: {val}','extractionMethod':'openpyxl-comprehensive-derived','statementContext':'analytical-input','confidence':1.0,'systemVerified':True,'validated':False,'rowBinding':'derived-from-explicit-workbook-inputs','verificationMethod':'workbook-derived-deterministic'})
    for fy in years:
        rev=by.get(('revenue',fy)); prior=by.get(('prior_revenue',fy))
        add_derived('beginning_value',fy,rev['normalizedValue'] if rev else None,'Beginning Value','Revenue')
        add_derived('ending_value',fy,rev['normalizedValue'] if rev else None,'Ending Value','Revenue')
        add_derived('number_of_years',fy,max(1,fy-min(years)),'Number of Years','Fiscal-year distance from earliest workbook year')
        da=by.get(('annual_dividend_per_share',fy)); shares=by.get(('weighted_avg_shares',fy))
        add_derived('dividends_paid',fy,(da['normalizedValue']*shares['normalizedValue']) if da and shares else None,'Dividends Paid','Annual dividend per share × weighted average shares')
        op=by.get(('operating_income',fy)); add_derived('ebit',fy,op['normalizedValue'] if op else None,'EBIT','Operating Income')
        churn=by.get(('churn_arr',fy)); add_derived('lost_revenue',fy,churn['normalizedValue'] if churn else None,'Lost Revenue','Churn ARR')
        mc=by.get(('market_cap',fy)); add_derived('market_value_equity',fy,mc['normalizedValue'] if mc else None,'Market Value Equity','Market Cap')
        mrr=by.get(('mrr',fy)); cust=by.get(('starting_customers',fy)); add_derived('mrr_per_customer',fy,(mrr['normalizedValue']/cust['normalizedValue']) if mrr and cust and cust['normalizedValue'] else None,'MRR Per Customer','MRR ÷ Starting Customers')
        add_derived('number_of_months',fy,12,'Number of Months','12 months per fiscal year')
        if prior and prior['normalizedValue']:
            growth=rev['normalizedValue']/prior['normalizedValue']-1 if rev else None
        else: growth=0.0 if rev else None
        add_derived('revenue_growth_rate',fy,growth,'Revenue Growth Rate','Revenue ÷ Prior Revenue − 1')
        add_derived('organic_growth',fy,growth,'Organic Growth Rate','Revenue growth proxy from workbook inputs')
        add_derived('starting_revenue',fy,prior['normalizedValue'] if prior else None,'Starting Revenue','Prior Revenue')
        taxes=by.get(('taxes',fy)); add_derived('tax_expense',fy,taxes['normalizedValue'] if taxes else None,'Tax Expense','Taxes')
        for concept in ('ebitda','operating_income'):
            prev=by.get((concept,fy-1)); add_derived('prior_'+concept,fy,prev['normalizedValue'] if prev else None,'Prior '+concept.replace('_',' ').title(),f'Prior-year {concept}')
    return facts

def run_generic(path, ext):
    if ext in ('.xlsx','.xls'):
        import openpyxl
        wb=openpyxl.load_workbook(path,data_only=True,read_only=True)
        pages=[]; tables=[]
        for s in wb.sheetnames:
            ws=wb[s]; rows=[]
            for row in ws.iter_rows(values_only=True):
                vals=[str(v) if v is not None else '' for v in row]
                if any(vals): rows.append('\t'.join(vals))
            pages.append({'page':len(pages)+1,'text':f'SHEET: {s}\n'+'\n'.join(rows),'extractor':'openpyxl'})
        facts=extract_comprehensive_workbook(path)
        return {'name':'openpyxl','method':'openpyxl','pages':pages,'tables':tables,'structuredFacts':facts,'documentFiscalYear':max(f.get('fiscalYear') for f in facts) if facts else None,'comparativeFiscalYears':sorted({int(f.get('fiscalYear')) for f in facts if f.get('fiscalYear')},reverse=True)}
    if ext=='.docx':
        from docx import Document
        d=Document(path); text='\n'.join(p.text for p in d.paragraphs)
        return {'name':'python-docx','method':'python-docx','pages':[{'page':1,'text':text,'extractor':'python-docx'}],'tables':[]}
    if ext in ('.html','.htm'):
        return run_html(path)
    return {'name':'native','method':'native','pages':[{'page':1,'text':Path(path).read_text(errors='ignore'),'extractor':'native'}],'tables':[]}



def score_fact(f):
    score=0.0
    ctx=f.get('statementContext')
    if ctx in ('income','balance','cashflow'): score+=0.35
    if f.get('sourcePage') is not None: score+=0.05
    method=str(f.get('extractionMethod','')).lower()
    if 'docling' in method: score+=0.20
    elif 'pdfplumber' in method: score+=0.15
    elif 'pymupdf' in method: score+=0.10
    if f.get('confidence') is not None: score+=float(f.get('confidence',0))*0.20
    ev=str(f.get('evidenceText','')).lower()
    if any(k in ev for k in ('total ','revenue from operations','profit for the year','net income','consolidated')): score+=0.20
    return score

def merge_fact_consensus(facts):
    groups={}
    for f in facts:
        key=(str(f.get('concept','')).lower(),int(f.get('fiscalYear') or 0),round(float(f.get('normalizedValue')),4) if f.get('normalizedValue') is not None else None)
        groups.setdefault(key,[]).append(f)
    by_concept_year={}
    for key,items in groups.items():
        concept,year,_=key
        consensus=len(items)
        best=max(items,key=score_fact)
        quality=min(1.0,0.55 + 0.15*consensus + 0.15*(1 if len({i.get('extractionMethod') for i in items})>1 else 0)+0.15*score_fact(best))
        best=dict(best)
        best['consensusCount']=consensus
        best['consensusExtractors']=sorted({i.get('extractionMethod') for i in items if i.get('extractionMethod')})
        best['consensusQuality']=round(quality,3)
        independent=len({str(i.get('extractionMethod')) for i in items if i.get('extractionMethod')}) >= 2
        strong_statement=all(str(i.get('statementContext','')) in {'income','balance','cashflow'} for i in items)
        metadata_consistent=(len({str(i.get('currency') or '').upper() for i in items})==1 and
                             len({str(i.get('scale') or '').lower() for i in items})==1 and
                             len({str(i.get('unit') or '').lower() for i in items})==1 and
                             len({str(i.get('fiscalYear') or '') for i in items})==1)
        aggregate_concepts={'revenue','cogs','gross_profit','operating_income','ebitda','net_income','cash','current_assets','assets','receivables','inventory','current_liabilities','payables','debt','current_debt','long_term_debt','liabilities','equity','operating_cash_flow','capex','depreciation_amortization','interest_expense'}
        deterministic_label_verifiable=(str(best.get('rowBinding',''))=='exact-label-row' and strong_statement and metadata_consistent and quality>=0.88 and concept in aggregate_concepts)
        best['systemVerified']=bool(deterministic_label_verifiable and ((consensus>=2 and quality>=0.88) or (consensus==1 and float(best.get('confidence') or 0)>=0.93)))
        if best['systemVerified'] and consensus==1:
            best['verificationMethod']='deterministic-statement-label'

        best['validated']=False
        best['verificationMethod']='multi-extractor-consensus' if best['systemVerified'] else 'source-extracted-needs-review'
        dest=by_concept_year.setdefault((concept,year),[])
        dest.append(best)
    out=[]
    for key,items in by_concept_year.items():
        # Prefer cross-engine agreement. If no agreement, use strongest statement/page signal.
        best=max(items,key=lambda f:(f.get('consensusCount',1),f.get('consensusQuality',0),score_fact(f)))
        out.append(best)
    return out

def derive_debt_aggregates(facts):
    by={}
    for f in facts:
        k=(str(f.get('fiscalYear') or ''), str(f.get('concept') or ''))
        by.setdefault(k,[]).append(f)
    years=sorted({str(f.get('fiscalYear') or '') for f in facts if f.get('fiscalYear')}, reverse=True)
    additions=[]
    for y in years:
        if any(str(f.get('fiscalYear') or '')==y and f.get('concept')=='debt' for f in facts): continue
        cur=next((f for f in facts if str(f.get('fiscalYear') or '')==y and f.get('concept')=='current_debt' and f.get('normalizedValue') is not None),None)
        lt=next((f for f in facts if str(f.get('fiscalYear') or '')==y and f.get('concept')=='long_term_debt' and f.get('normalizedValue') is not None),None)
        if cur and lt:
            total=float(cur['normalizedValue'])+float(lt['normalizedValue'])
            additions.append({'concept':'debt','sourceLabel':'Debt (derived from current and non-current debt)','aggregateRole':'derived-aggregate','rawValue':str(total),'normalizedValue':total,'absoluteValue':total,'baseValue':total,'unit':cur.get('unit') or lt.get('unit'),'currency':cur.get('currency') or lt.get('currency'),'scale':cur.get('scale') or lt.get('scale') or 'units','sourceUnitText':cur.get('sourceUnitText') or lt.get('sourceUnitText'),'fiscalYear':y,'sourcePage':cur.get('sourcePage') or lt.get('sourcePage'),'evidenceText':f"Derived debt = {cur.get('rawValue')} + {lt.get('rawValue')}",'extractionMethod':'deterministic-aggregate','statementContext':'balance','confidence':0.99,'systemVerified':False,'validated':False,'rowBinding':'derived-aggregate','derivedFromFacts':[cur.get('id'),lt.get('id')]})
    return facts+additions

def extract(path):
    path=str(path); ext=Path(path).suffix.lower()
    engines=[]
    if ext in ('.html','.htm'):
        r=run_html(path)
        facts=derive_debt_aggregates(merge_fact_consensus(r.get('structuredFacts',[])))
        years=sorted({int(f.get('fiscalYear')) for f in facts if f.get('fiscalYear')}, reverse=True)
        cur,scale,unit=detect_money_metadata(r.get('pages',[{'text':''}])[0].get('text',''))
        for f in facts:
            f['currency']=f.get('currency') or cur; f['scale']=f.get('scale') if f.get('scale')!='units' else (scale or 'units'); f['unit']=f.get('unit') or unit or 'document unit'
        return {'text':r['pages'][0]['text'],'pages':1,'method':'html-table-structured','tables':r.get('tables',[]),'images':r.get('images',[]),'comparativeFiscalYears':years,'primaryExtractor':r['name'],'extractors':[{'name':r['name'],'method':r['method'],'ok':True,'pageCount':1,'tableCount':len(r.get('tables',[]))}], 'documentFiscalYear':filename_document_year(path) or detect_document_year(r['pages'][0]['text'],path), 'documentUnit':unit or (f'{cur} {scale}' if cur and scale else cur), 'documentCurrency':cur,'documentScale':scale or 'units','structuredFacts':facts,'extractionQuality':{'grade':'high' if facts else ('rich' if r.get('tables') else 'text'),'comparativeFiscalYears':years,'comparativeColumnCount':len(years),'tableCount':len(r.get('tables',[])),'imageCount':len(r.get('images',[])),'financialDocument':True,'factCount':len(facts),'ragFeedExtractor':'html-table-structured'}}
    if ext=='.pdf':
        # Cheap independent extractors run in parallel. Heavy layout engines are
        # escalated only when the cheap pass signals that they are likely useful.
        jobs=[run_pymupdf]
        try:
            import pdfplumber  # noqa: F401
            jobs.append(run_pdfplumber)
        except Exception: pass
        results=[]
        with ThreadPoolExecutor(max_workers=min(MAX_WORKERS,len(jobs))) as ex:
            futs=[ex.submit(j,path) for j in jobs]
            for fut in as_completed(futs):
                try: results.append(fut.result())
                except Exception as e: results.append({'name':getattr(fut,'__name__','engine'),'error':str(e),'pages':[],'tables':[]})
        text='\n\n'.join(p.get('text','') for r in results for p in r.get('pages',[]))
        # pypdf is an independent fallback, not an always-on peer: keep the stable
        # born-digital extraction path authoritative when it already has useful text.
        sparse_primary=(len(re.sub(r'\s+','',text))<800 or len([p for r in results for p in r.get('pages',[]) if p.get('text')])<2)
        if sparse_primary:
            try:
                import pypdf  # noqa: F401
                pp=run_pypdf(path)
                pp_text='\n\n'.join(p.get('text','') for p in pp.get('pages',[]))
                if len(re.sub(r'\s+','',pp_text))>len(re.sub(r'\s+','',text))*0.75:
                    results.append(pp)
                    text='\n\n'.join(p.get('text','') for r in results for p in r.get('pages',[]))
            except Exception as e:
                results.append({'name':'pypdf','method':'pypdf','error':str(e),'pages':[],'tables':[]})
        # For financial PDFs, run at least one structure-aware engine alongside the cheap text engines.
        # Resource control is explicit: cheap engines share MAX_WORKERS; heavy engines run with
        # a separate bounded pool so multiple extractors can be parallel without spawning an
        # unbounded OCR/layout worker storm. Additional heavy engines are only added when installed.
        financial=bool(re.search(r'(balance sheet|profit and loss|statement of profit|cash flows|revenue from operations|consolidated financial)', text, re.I))
        # OCR is an explicit fallback, not the default path: invoke it only when the combined text layer is too sparse for a document-sized PDF.
        ocr_result=None
        sparse_text=(len(re.sub(r'\s+','',text))<800 or len([p for r in results for p in r.get('pages',[]) if p.get('text')])<2)
        if sparse_text and financial and os.getenv('MYAI_CFO_ENABLE_OCR','1')!='0':
            try:
                ocr_result=run_tesseract_ocr(path)
                if ocr_result.get('pages'):
                    results.append(ocr_result)
                    ocr_text='\n\n'.join(p.get('text','') for p in ocr_result.get('pages',[]))
                    if len(ocr_text)>len(text): text=ocr_text
            except Exception as e:
                results.append({'name':'Tesseract OCR','method':'tesseract-ocr','error':str(e),'pages':[],'tables':[]})
        if financial:
            heavy=[]
            heavy_enabled=os.getenv('MYAI_CFO_ENABLE_HEAVY_EXTRACTORS','')=='1'
            # Keep the financial evidence spine deterministic and resource-bounded by
            # default on every runtime, not only certification. Optional layout-model
            # engines (Docling/Marker/MinerU/NVIDIA) are explicitly opt-in. This
            # prevents a later document from destabilising an earlier committed
            # extraction while preserving the advanced engines for deliberate use.
            financial_heavy_disabled=not heavy_enabled
            if not financial_heavy_disabled:
                try:
                    import nv_ingest_client  # noqa: F401
                    heavy.append(('NVIDIA NeMo Retriever',run_nvidia_retriever))
                except Exception: pass
                try:
                    import docling  # noqa: F401
                    heavy.append(('Docling',run_docling))
                except Exception: pass
                import shutil
                if shutil.which('marker_single'): heavy.append(('Marker',run_marker))
                if shutil.which('mineru'): heavy.append(('MinerU',run_mineru))
            heavy_workers=max(1,min(int(os.getenv('MYAI_HEAVY_EXTRACTOR_WORKERS','1')),len(heavy))) if heavy else 0
            if heavy_workers:
                with ThreadPoolExecutor(max_workers=heavy_workers) as ex:
                    futs={ex.submit(fn,path):name for name,fn in heavy}
                    for fut in as_completed(futs):
                        name=futs[fut]
                        try: results.append(fut.result())
                        except Exception as e: results.append({'name':name,'error':str(e),'pages':[],'tables':[]})
        # Optional PdfTable is an additional independent table signal. It is not mandatory
        # and is only attempted when installed (or when explicitly requested) to avoid
        # adding heavyweight model startup to every PDF ingestion.
        if financial and os.getenv('MYAI_CFO_ENABLE_PDFTABLE','1')!='0':
            try:
                pt=run_pdftable(path)
                results.append(pt)
                if pt.get('pages'):
                    text=text+'\n\n'+"\n\n".join(p.get('text','') for p in pt.get('pages',[]))
            except Exception as e:
                results.append({'name':'PdfTable','method':'pdftable','pages':[],'tables':[],'error':str(e),'optional':True})
        # Excel certification workbooks often place comparative years in the first header row.
        # The latest reported year is the document period; never let the first column year
        # (e.g. 2023) become the document fiscal year when the workbook contains 2023/2024/2025.
        excel_years=[]
        if Path(path).suffix.lower() in ('.xlsx','.xls'):
            for m in re.finditer(r'(?<!\d)(20\d{2})(?!\d)', text[:20000]): excel_years.append(int(m.group(1)))
            excel_years=sorted(set(excel_years))
        doc_year=filename_document_year(path) or (max(excel_years) if len(excel_years)>=2 else None) or detect_document_year(text, path)
        # Prefer page-order from the first structured extractor and limit to the first
        # parent statement set to avoid subsidiary statements near the end of group reports.
        base_pages=[]
        # Prefer a layout-aware born-digital text pass when available; its inline
        # table rows preserve comparative values better than line-separated PDF text.
        # Prefer the layout-aware engine for the RAG feed when installed, while
        # retaining independent lightweight engines for fact reconciliation.
        ordered_results=sorted(results,key=lambda r: {'NVIDIA NeMo Retriever':0,'Docling':1,'pdfplumber':2,'PdfTable':3,'PyMuPDF':4,'pypdf':5,'Tesseract OCR':6}.get(r.get('name'),9))
        for r in ordered_results:
            if r.get('pages'): base_pages=r['pages']; break
        selected=primary_statement_pages(base_pages)
        if not selected: selected=base_pages[:16]
        facts=[]
        # Extract against every available engine's candidate pages, then consensus-merge.
        # This prevents a single parser's narrative row from becoming the authoritative fact.
        for r in results:
            ps=primary_statement_pages(r.get('pages',[])) or r.get('pages',[])[:16]
            for p in ps: facts.extend(facts_from_page(p,doc_year,path))
        facts=derive_debt_aggregates(merge_fact_consensus(facts))
        document_currency, document_scale, document_unit = detect_money_metadata(text)
        conflicts=[]
        bycy={}
        for f in facts:
            bycy.setdefault((f.get('concept'),f.get('fiscalYear')),[]).append(f)
        for key,items in bycy.items():
            vals={round(float(i.get('normalizedValue')),4) for i in items if i.get('normalizedValue') is not None}
            if len(vals)>1: conflicts.append({'concept':key[0],'fiscalYear':key[1],'values':sorted(vals),'count':len(vals)})
        for f in facts:
            f.setdefault('currency', document_currency)
            f.setdefault('scale', document_scale or 'units')
            f.setdefault('unit', document_unit or f.get('unit') or 'document unit')
        verified=sum(1 for f in facts if f.get('systemVerified'))
        table_records=[]
        seen_tables=set()
        for r in results:
            for t in r.get('tables',[]) or []:
                try:
                    if isinstance(t, list):
                        rows=['\t'.join('' if c is None else str(c) for c in row) for row in t]
                        rec={'title':'Detected table','page':None,'tsv':'\n'.join(rows)}
                    else:
                        rec={'title':t.get('title') or 'Detected table','page':t.get('page'),'tsv':t.get('tsv') or ''}
                    key=json.dumps(rec,sort_keys=True,ensure_ascii=False)
                    if key not in seen_tables and rec.get('tsv'):
                        seen_tables.add(key); table_records.append(rec)
                except Exception:
                    pass
        comparative_years=sorted({int(f.get('fiscalYear')) for f in facts if f.get('fiscalYear')},reverse=True)
        return {
            'text': text,
            'pages': len(base_pages),
            'method':'ensemble',
            'tables': table_records,
            'comparativeFiscalYears':sorted(set(comparative_years + (excel_years if 'excel_years' in locals() else [])),reverse=True),
            'primaryExtractor': ordered_results[0].get('name') if ordered_results else 'none',
            'extractors':[{'name':r.get('name'), 'method':r.get('method'), 'ok':bool(r.get('pages')),'pageCount':len(r.get('pages',[])),'tableCount':len(r.get('tables',[])),'ocrPages':r.get('ocrPages',0),'error':r.get('error')} for r in results],
            'resourcePolicy':{'maxWorkers':MAX_WORKERS,'heavyLayoutEscalation':'bounded-parallel','heavyWorkers':max(0,min(int(os.getenv('MYAI_HEAVY_EXTRACTOR_WORKERS','2')),3))},
            'documentFiscalYear':doc_year,
            'documentUnit':document_unit,
            'documentCurrency':document_currency,
            'documentScale':document_scale or 'units',
            'structuredFacts':facts,
            'extractionQuality':{
                'grade':'ensemble' if facts else 'needs-review',
                'comparativeFiscalYears':comparative_years,
                'comparativeColumnCount':len(comparative_years),
                'tableCount':len(table_records),
                'financialDocument':financial,
                'factCount':len(facts),
                'systemVerifiedFactCount':verified,
                'statementPages':len(selected),
                'engineCount':len(results),
                'engineOrder':[r.get('name') for r in ordered_results if r.get('pages')],
                'ragFeedExtractor':ordered_results[0].get('name') if ordered_results else 'none',
                'conflicts':conflicts,
                'nvidiaUsed':any(r.get('name')=='NVIDIA NeMo Retriever' and r.get('pages') for r in results),
                'doclingUsed':any(r.get('name')=='Docling' and r.get('pages') for r in results)
            }
        }
    r=run_generic(path,ext)
    text='\n'.join(p.get('text','') for p in r.get('pages',[]))
    excel_years=[]
    if ext in ('.xlsx','.xls'):
        excel_years=sorted({int(y) for y in re.findall(r'(?<!\d)(20\d{2})(?!\d)', '\n'.join(text.splitlines()[:12])) if 2010<=int(y)<=2100})
    doc_year=filename_document_year(path) or (max(excel_years) if len(excel_years)>=2 else None) or detect_document_year(text, path)
    if ext in ('.xlsx','.xls'):
        r['comparativeFiscalYears']=excel_years[-4:][::-1]
        r['extractionQuality']={**(r.get('extractionQuality') or {}),'comparativeFiscalYears':r['comparativeFiscalYears'],'comparativeColumnCount':len(r['comparativeFiscalYears'])}
    facts=list(r.get('structuredFacts',[])) if ext in ('.xlsx','.xls') and r.get('structuredFacts') else []
    if not facts:
        for p in r.get('pages',[]): facts.extend(facts_from_page(p,doc_year,path))
    dedup={}
    for f in facts:
        dedup[(f['concept'],f.get('fiscalYear'))]=f
    return {**r,'text':text,'documentFiscalYear':doc_year,'structuredFacts':list(dedup.values()),'extractionQuality':{'grade':'structured','factCount':len(dedup)}}


def main():
    ap=argparse.ArgumentParser(); ap.add_argument('--input',required=True); ap.add_argument('--output',required=True); ap.add_argument('--assets',default='')
    a=ap.parse_args(); result=extract(a.input); Path(a.output).write_text(json.dumps(result,ensure_ascii=False,indent=2),encoding='utf-8')
    # Always emit compact JSON on stdout for server diagnostics.
    print(json.dumps({'ok':True,'method':result.get('method'),'factCount':len(result.get('structuredFacts',[])),'documentFiscalYear':result.get('documentFiscalYear'),'primaryExtractor':result.get('primaryExtractor'),'extractors':result.get('extractors',[])},ensure_ascii=False))

if __name__=='__main__': main()
